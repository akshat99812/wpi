#!/usr/bin/env python3
"""inventory_docx.py — Phase 0 step 2 inventory.

Walks every .docx under ingestion/data/docx/ and emits
ingestion/data/docx_inventory.json with: filename, year, size_bytes,
paragraph_count, heading_count, table_count, table_row_total,
largest_table_rows, image_count, est_pages, matching_pdf.

Read-only. No side effects beyond writing the JSON.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import docx
from docx.document import Document as DocxDocument

ROOT = Path(__file__).resolve().parent
DOCX_DIR = ROOT / "data" / "docx"
PDF_DIR = ROOT / "data" / "pdfs"
OUT = ROOT / "data" / "docx_inventory.json"


def infer_year(name: str) -> int | None:
    m = re.search(r"(19|20)\d{2}", name)
    return int(m.group(0)) if m else None


def matching_pdf(year: int | None) -> str | None:
    if year is None or not PDF_DIR.exists():
        return None
    for p in PDF_DIR.rglob("*.pdf"):
        if str(year) in p.name:
            return str(p.relative_to(ROOT))
    return None


def analyse(path: Path) -> dict:
    d: DocxDocument = docx.Document(str(path))
    paragraphs = list(d.paragraphs)
    tables = list(d.tables)
    headings = [p for p in paragraphs if (p.style and p.style.name and p.style.name.startswith("Heading"))]
    table_row_total = sum(len(t.rows) for t in tables)
    largest_table_rows = max((len(t.rows) for t in tables), default=0)
    body_xml = d.element.body.xml
    image_count = body_xml.count("<w:drawing")
    chars = sum(len(p.text) for p in paragraphs) + sum(
        sum(len(c.text) for row in t.rows for c in row.cells) for t in tables
    )
    est_pages = max(1, round(chars / 1800))
    year = infer_year(path.name)
    return {
        "filename": path.name,
        "relpath": str(path.relative_to(ROOT)),
        "year": year,
        "size_bytes": path.stat().st_size,
        "paragraph_count": len(paragraphs),
        "heading_count": len(headings),
        "table_count": len(tables),
        "table_row_total": table_row_total,
        "largest_table_rows": largest_table_rows,
        "image_count": image_count,
        "est_pages": est_pages,
        "char_count": chars,
        "matching_pdf": matching_pdf(year),
    }


def main() -> None:
    files = sorted(DOCX_DIR.rglob("*.docx"))
    if not files:
        sys.exit(f"FATAL: no .docx files under {DOCX_DIR}")
    report = []
    for f in files:
        try:
            report.append(analyse(f))
            print(f"  OK  {f.name}")
        except Exception as e:
            print(f"  ERR {f.name}: {e}", file=sys.stderr)
            report.append({"filename": f.name, "error": str(e)})
    OUT.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(f"\nwrote {OUT.relative_to(ROOT.parent)} ({len(report)} entries)")


if __name__ == "__main__":
    main()

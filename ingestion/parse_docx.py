#!/usr/bin/env python3
"""parse_docx.py — Phase 1 parser. Lossless docx → markdown + elements.jsonl.

This is the most important file in the rebuild. If parsing drops or
corrupts a table cell, no amount of clever chunking or retrieval tuning
will recover it. Read carefully before changing.

Behaviour:
    - Walks document.element.body in document order so paragraphs/tables
      stay interleaved (python-docx's `document.paragraphs` + `document.tables`
      flattens this and loses sequence — do not use them here).
    - Headings update a section-path stack. In this corpus most docs use
      `Normal` style throughout (no real heading structure), so we ALSO
      seed the stack at level 0 with a section name derived from the
      filename (e.g. "STATISTICS" from "FIN SEC 6 - 2025 - STATISTICS").
    - Tables are emitted with full cell data, an explicit header row, an
      auto-generated natural-language caption, AND `title_context`: the
      1-3 non-empty paragraphs immediately preceding the table (the
      table's visual title block). The caption embeds the title_context
      so downstream chunks carry it.
    - Merged cells: python-docx exposes the SAME text in each merged cell.
      We KEEP that duplication so no data is silently dropped, and flag
      the merge in the parse report so the chunker/eval can see it.
    - Images are detected and reported, never OCR'd at this stage. Text
      boxes (`w:txbxContent`) are DROPPED — in this corpus they hold
      chart-axis ticks and bar-data labels ("77.700 MW", "0", "20", ...)
      with no structural link to which chart they belong to. Embedding
      them as standalone fragments produces noise without answer value.

Usage:
    python parse_docx.py <docx_path> [--out-dir data/parsed]

Outputs (alongside each other in --out-dir):
    <stem>.md                    — pretty GitHub-flavored markdown
    <stem>.elements.jsonl        — one element per line (heading | paragraph | table | image | textbox)
    <stem>.parse_report.json     — warnings, anomalies, image/textbox list, counts
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable, _Cell
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
DEFAULT_OUT = ROOT / "data" / "parsed"


# ────────────────────────────────────────────────────────────────────────────
# Heading-stack helpers
# ────────────────────────────────────────────────────────────────────────────

@dataclass
class SectionStack:
    """Tracks current heading path. Indexed by heading level.

    A Heading-N replaces the entries at depth >= N. So a Heading-2 always
    has its Heading-1 ancestor still on the stack (if one existed)."""
    levels: dict[int, str] = field(default_factory=dict)

    def push(self, level: int, text: str) -> None:
        # Drop any deeper headings — they belong to the previous section.
        for k in list(self.levels.keys()):
            if k >= level:
                del self.levels[k]
        self.levels[level] = text

    def path(self) -> list[str]:
        return [self.levels[k] for k in sorted(self.levels.keys())]


# ────────────────────────────────────────────────────────────────────────────
# Paragraph / heading
# ────────────────────────────────────────────────────────────────────────────

HEADING_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
YEAR_IN_NAME_RE = re.compile(r"[-_\s]*(19|20)\d{2}[-_\s]*")
FILENAME_PREFIX_RE = re.compile(r"^(FIN\s+)?(SEC\s+\d+\s*-?\s*)?", re.IGNORECASE)


def heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    m = HEADING_RE.match(style_name.strip())
    return int(m.group(1)) if m else None


def section_from_filename(stem: str) -> str:
    """Derive a clean section name from a docx filename.

    Examples:
        'FIN SEC 6 - 2025 - STATISTICS'         -> 'STATISTICS'
        'FIN SEC 11 - 2025 - OTHEWR INSTI. & PUB' -> 'OTHEWR INSTI. & PUB'
        'FIN OPENING PAGES-2025'                -> 'OPENING PAGES'
        'FIN CONTENT - 2025'                    -> 'CONTENT'
    """
    s = YEAR_IN_NAME_RE.sub(" ", stem).strip()
    s = FILENAME_PREFIX_RE.sub("", s).strip()
    s = re.sub(r"\s*-\s*", " - ", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip(" -") or stem


def clean_text(s: str) -> str:
    """Collapse whitespace within a single run of text, but preserve newlines
    as explicit \\n so a paragraph that spans multiple runs reads naturally."""
    # python-docx already joins runs into paragraph.text; just normalise spaces.
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse runs of spaces/tabs (but not newlines).
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


# ────────────────────────────────────────────────────────────────────────────
# Table extraction
# ────────────────────────────────────────────────────────────────────────────

def cell_plain_text(cell: _Cell) -> str:
    """Cell text with paragraph breaks preserved as ' / '. Strips drawings.

    Pipes inside cell text are escaped so they don't break markdown rendering
    when we emit the table later. JSON output keeps the raw text.
    """
    parts = [p.text for p in cell.paragraphs]
    text = " / ".join(s.strip() for s in parts if s and s.strip())
    return clean_text(text)


def cell_tc_id(cell: _Cell) -> int:
    """Identity of the underlying <w:tc>. Two python-docx cells that point
    to the same <w:tc> are the SAME merged cell."""
    return id(cell._tc)


def is_explicit_header_row(row) -> bool:
    """<w:trPr>/<w:tblHeader> marks a row as a repeating header."""
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is None:
        return False
    return trPr.find(qn("w:tblHeader")) is not None


def table_to_dict(t: DocxTable, section_path: list[str], idx: int, title_context: list[str]) -> tuple[dict, list[str]]:
    """Returns (table_element, warnings).

    `title_context` is the list of 1-3 preceding paragraph texts (the visual
    title block above the table — these docx files don't use Heading styles
    so we synthesise context from preceding paragraphs). It's embedded into
    the caption so downstream chunks carry the table's title."""
    warnings: list[str] = []
    raw_rows: list[list[str]] = []
    merge_map: list[list[int]] = []  # parallel to raw_rows, holds tc_id

    for row in t.rows:
        row_cells: list[str] = []
        row_ids: list[int] = []
        for cell in row.cells:
            row_cells.append(cell_plain_text(cell))
            row_ids.append(cell_tc_id(cell))
        raw_rows.append(row_cells)
        merge_map.append(row_ids)

    # Detect merged cells (same tc_id repeated in a row OR across rows).
    merged_horiz = 0
    merged_vert = 0
    for r, ids in enumerate(merge_map):
        for c in range(1, len(ids)):
            if ids[c] == ids[c - 1]:
                merged_horiz += 1
        if r > 0 and len(merge_map[r - 1]) == len(ids):
            for c in range(len(ids)):
                if ids[c] == merge_map[r - 1][c]:
                    merged_vert += 1
    if merged_horiz:
        warnings.append(f"table#{idx}: {merged_horiz} horizontally-merged cells (duplication preserved)")
    if merged_vert:
        warnings.append(f"table#{idx}: {merged_vert} vertically-merged cells (duplication preserved)")

    # Header row: explicit <w:tblHeader> on any row, else first row by default.
    header_idx = 0
    for i, row in enumerate(t.rows):
        if is_explicit_header_row(row):
            header_idx = i
            break

    header_row = raw_rows[header_idx] if raw_rows else []
    data_rows = raw_rows[header_idx + 1:] if raw_rows else []

    caption = make_caption(header_row, data_rows, section_path, idx, title_context)

    # Width sanity: ragged rows hint at parser confusion or merged-cell artefacts.
    widths = {len(r) for r in raw_rows}
    if len(widths) > 1:
        warnings.append(f"table#{idx}: ragged row widths {sorted(widths)}")

    element = {
        "type": "table",
        "index": idx,
        "section_path": section_path,
        "title_context": title_context,
        "header_row": header_row,
        "rows": raw_rows,         # includes header row at position header_idx
        "header_row_index": header_idx,
        "row_count": len(raw_rows),
        "data_row_count": len(data_rows),
        "column_count": max(widths) if widths else 0,
        "caption": caption,
    }
    return element, warnings


YEAR_RE = re.compile(r"(19|20)\d{2}")


def make_caption(header: list[str], data_rows: list[list[str]], section_path: list[str], idx: int, title_context: list[str]) -> str:
    section = section_path[-1] if section_path else "(no section)"
    cols = [h for h in header if h]
    cols_str = ", ".join(cols) if cols else "(unlabelled columns)"
    n = len(data_rows)
    coverage = ""
    if data_rows:
        first_col = [row[0] for row in data_rows if row]
        years: list[int] = []
        for v in first_col:
            for m in YEAR_RE.finditer(str(v)):
                years.append(int(m.group(0)))
        if len(set(years)) >= 2:
            coverage = f"; appears to cover years {min(years)}-{max(years)}"
    title_prefix = ""
    if title_context:
        # Join with " " so multi-line title blocks read as one sentence.
        title_prefix = " ".join(title_context).strip()
        title_prefix = re.sub(r"\s+", " ", title_prefix)
        title_prefix = f"'{title_prefix}'. "
    return f"{title_prefix}Table {idx} in section '{section}': columns are {cols_str}; {n} data rows{coverage}"


def render_table_md(header: list[str], rows: list[list[str]], header_idx: int) -> str:
    """GitHub-flavored pipe table. Pipes inside cells are escaped."""
    def esc(s: str) -> str:
        return s.replace("|", "\\|").replace("\n", " ")
    width = max((len(r) for r in rows), default=len(header))
    def pad(r: list[str]) -> list[str]:
        return list(r) + [""] * (width - len(r))
    out: list[str] = []
    hdr = pad(header)
    out.append("| " + " | ".join(esc(c) for c in hdr) + " |")
    out.append("|" + "|".join(["---"] * width) + "|")
    for i, r in enumerate(rows):
        if i == header_idx:
            continue
        out.append("| " + " | ".join(esc(c) for c in pad(r)) + " |")
    return "\n".join(out)


# ────────────────────────────────────────────────────────────────────────────
# Images
# ────────────────────────────────────────────────────────────────────────────

def drawing_alt_text(drawing_el) -> str:
    """Pull alt-text / description from a <w:drawing>. Empty string if none."""
    for docPr in drawing_el.iter(qn("wp:docPr")):
        descr = docPr.get("descr") or docPr.get("title") or docPr.get("name")
        if descr:
            return descr
    return ""


def collect_drawings_in_paragraph(p: Paragraph, section_path: list[str]) -> list[dict]:
    """Find <w:drawing> nodes anywhere under the paragraph element."""
    images: list[dict] = []
    for drawing in p._p.iter(qn("w:drawing")):
        images.append({
            "type": "image",
            "alt_text": drawing_alt_text(drawing) or "no alt-text",
            "section_path": section_path,
        })
    return images


# Text boxes intentionally NOT extracted. In this corpus <w:txbxContent>
# holds chart-axis ticks and bar-data labels ("77.700 MW", "0", "20", ...)
# with no link to which chart they belong to. Embedding them produces
# noise without answer value. If a future corpus puts narrative text in
# text boxes, restore extraction here.


# ────────────────────────────────────────────────────────────────────────────
# Main walk
# ────────────────────────────────────────────────────────────────────────────

def parse_doc(docx_path: Path, out_dir: Path) -> dict:
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = docx_path.stem
    md_path = out_dir / f"{stem}.md"
    elements_path = out_dir / f"{stem}.elements.jsonl"
    report_path = out_dir / f"{stem}.parse_report.json"

    d: DocxDocument = docx.Document(str(docx_path))
    stack = SectionStack()

    # Seed the section stack at level 0 with a filename-derived section.
    # In this corpus most paragraphs are styled 'Normal' so there are no
    # real headings; the filename ("FIN SEC 6 - 2025 - STATISTICS") carries
    # the only reliable top-level section signal.
    filename_section = section_from_filename(docx_path.stem)
    stack.push(0, filename_section)

    elements: list[dict] = []
    md_lines: list[str] = []
    warnings: list[str] = []
    table_idx = 0
    image_locations: list[dict] = []

    # Rolling buffer of the most recent non-empty, non-heading paragraphs.
    # When we hit a table, the last N entries form its `title_context`.
    # Reset on heading or table — title context must be contiguous.
    # Paragraphs that look like POSTSCRIPTS of the previous table
    # ("Note: ...", "(Source: MNRE)", "* ..." footnotes) are still emitted
    # as regular paragraph elements but EXCLUDED from the buffer so they
    # don't bleed into the next table's title.
    title_buf: list[str] = []
    MAX_TITLE_CONTEXT = 5  # title blocks here run up to 4 lines, +1 slack
    POSTSCRIPT_RE = re.compile(r"^\s*(note\s*[:\-]|source\s*[:\-]|\(\s*source|\*\s|@\s|see\s+also)", re.IGNORECASE)
    # Boilerplate that appears between split tables: "Contd……", "Contd...",
    # "Continued", "(Contd.)" — dropped at parse time so it never reaches
    # the chunker as a 23-token noise chunk.
    BOILERPLATE_RE = re.compile(r"^\s*\(?\s*cont(d|inue|inued)?\b[\.\…]*\s*\)?\s*$", re.IGNORECASE)

    md_lines.append(f"# {filename_section} (year unknown from name)" if not re.search(r"(19|20)\d{2}", docx_path.stem) else f"# {filename_section} ({re.search(r'(19|20)\d{2}', docx_path.stem).group(0)})")
    md_lines.append("")

    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, d._body)
            text = clean_text(p.text)
            style = p.style.name if p.style else None
            lvl = heading_level(style)

            # Capture any inline images (drawings) regardless of paragraph type.
            for img in collect_drawings_in_paragraph(p, stack.path()):
                elements.append(img)
                image_locations.append({
                    "alt_text": img["alt_text"],
                    "section_path": img["section_path"],
                })
                md_lines.append(f"[IMAGE: {img['alt_text']}]")
                md_lines.append("")

            if lvl is not None and text:
                stack.push(lvl, text)
                elements.append({
                    "type": "heading",
                    "level": lvl,
                    "text": text,
                    "section_path": stack.path(),
                })
                md_lines.append(f"{'#' * min(lvl, 6)} {text}")
                md_lines.append("")
                title_buf.clear()  # headings reset the title context
            elif text:
                if BOILERPLATE_RE.match(text):
                    # "Contd…" / "Continued" between split tables — drop entirely.
                    continue
                # Drop pure-punctuation / single-char paragraphs (e.g. stray
                # "(" or ")" emitted by python-docx around drawings/shapes).
                if len(re.sub(r"[\s\W_]+", "", text)) < 2:
                    continue
                elements.append({
                    "type": "paragraph",
                    "text": text,
                    "section_path": stack.path(),
                })
                md_lines.append(text)
                md_lines.append("")
                if POSTSCRIPT_RE.match(text):
                    # Likely a note/source postscript of the prior table;
                    # don't let it pollute the next table's title context.
                    title_buf.clear()
                else:
                    title_buf.append(text)
                    if len(title_buf) > MAX_TITLE_CONTEXT:
                        title_buf.pop(0)
            # Empty paragraphs are dropped — they carry no information.

        elif child.tag == qn("w:tbl"):
            t = DocxTable(child, d._body)
            table_idx += 1
            elt, w = table_to_dict(t, stack.path(), table_idx, list(title_buf))
            warnings.extend(w)
            elements.append(elt)
            md_lines.append(f"<!-- {elt['caption']} -->")
            md_lines.append(render_table_md(elt["header_row"], elt["rows"], elt["header_row_index"]))
            md_lines.append("")
            title_buf.clear()  # tables consume the preceding title block

        # Other body children (sectPr, etc.) are ignored — they carry no
        # user-facing content for our RAG purposes.

    # ── write outputs ─────────────────────────────────────────────────────
    md_path.write_text("\n".join(md_lines).rstrip() + "\n", encoding="utf-8")
    with elements_path.open("w", encoding="utf-8") as fh:
        for el in elements:
            fh.write(json.dumps(el, ensure_ascii=False) + "\n")

    counts = {
        "headings": sum(1 for e in elements if e["type"] == "heading"),
        "paragraphs": sum(1 for e in elements if e["type"] == "paragraph"),
        "tables": sum(1 for e in elements if e["type"] == "table"),
        "images": sum(1 for e in elements if e["type"] == "image"),
    }
    table_stats = []
    for e in elements:
        if e["type"] == "table":
            table_stats.append({
                "index": e["index"],
                "section": e["section_path"][-1] if e["section_path"] else None,
                "rows": e["row_count"],
                "data_rows": e["data_row_count"],
                "columns": e["column_count"],
                "caption": e["caption"],
            })
    report = {
        "source": str(docx_path.relative_to(ROOT) if docx_path.is_relative_to(ROOT) else docx_path),
        "source_sha1": hashlib.sha1(docx_path.read_bytes()).hexdigest(),
        "filename_section": filename_section,
        "counts": counts,
        "warnings": warnings,
        "images": image_locations,
        "tables": table_stats,
        "outputs": {
            "md": str(md_path.relative_to(ROOT)),
            "elements": str(elements_path.relative_to(ROOT)),
        },
    }
    report_path.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    return report


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path)
    ap.add_argument("--out-dir", type=Path, default=DEFAULT_OUT)
    args = ap.parse_args()
    if not args.docx.exists():
        sys.exit(f"FATAL: not found: {args.docx}")
    report = parse_doc(args.docx, args.out_dir)
    print(json.dumps(report["counts"], indent=2))
    if report["warnings"]:
        print(f"\n{len(report['warnings'])} warning(s):")
        for w in report["warnings"]:
            print(f"  - {w}")
    print(f"\nmd:       {report['outputs']['md']}")
    print(f"elements: {report['outputs']['elements']}")


if __name__ == "__main__":
    main()
